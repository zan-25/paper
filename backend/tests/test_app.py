import os
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

os.environ["DATABASE_URL"] = "sqlite://"
os.environ["STORAGE_ROOT"] = "D:/Projects/QPGen/Front-End-Builder/backend/test-storage"

Path("D:/tmp").mkdir(parents=True, exist_ok=True)
Path("D:/Projects/QPGen/Front-End-Builder/backend/test-storage").mkdir(parents=True, exist_ok=True)

from app.main import app

def login(client: TestClient, email: str, password: str) -> str:
    response = client.post("/api/v1/auth/login", json={"email": email, "password": password})
    assert response.status_code == 200
    return response.json()["access_token"]


def _question_table(document: DocxDocument):
    return next(
        table
        for table in document.tables
        if table.rows
        and table.rows[0].cells
        and table.rows[0].cells[0].text.strip() == "Q No"
    )


def test_teacher_to_hod_review_flow() -> None:
    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        generate_response = client.post(
            "/api/v1/papers/generate",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": 1,
                "title": "IAT-1 Question Paper",
                "exam_type": "IAT-1",
                "semester": "5",
                "batch": "2022-26",
                "max_marks": 30,
                "duration_minutes": 90,
                "teaching_department": "AIML",
                "prompt": "Generate a balanced paper for CO1-CO3 with a mix of L1-L4.",
                "rbt_levels": ["L1", "L2", "L3", "L4"],
                "module_numbers": [1, 2, 3, 4],
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()

        submit_response = client.post(
            f"/api/v1/papers/{paper['id']}/submit",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert submit_response.status_code == 200
        assert submit_response.json()["status"] == "pending_review"

        hod_token = login(client, "hod@dsatm.edu", "Hod@123")
        review_response = client.post(
            f"/api/v1/reviews/{paper['id']}/action",
            headers={"Authorization": f"Bearer {hod_token}"},
            json={"decision": "approved", "comments": "Balanced and ready for exam cell."},
        )
        assert review_response.status_code == 200
        assert review_response.json()["status"] == "approved"

        download_response = client.get(
            f"/api/v1/papers/{paper['id']}/download",
            headers={"Authorization": f"Bearer {hod_token}"},
        )
        assert download_response.status_code == 200
        exported = DocxDocument(BytesIO(download_response.content))
        paragraph_text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
        table_text = "\n".join(
            cell.text for table in exported.tables for row in table.rows for cell in row.cells
        )
        assert "Dayananda Sagar Academy of Technology & Management" in table_text
        assert "USN:" in table_text
        assert "Department of Artificial Intelligence and Machine Learning" in paragraph_text or table_text
        assert "Percentage of CO Coverage" in paragraph_text
        assert "Percentage of Syllabus coverage" in paragraph_text
        assert "COs" in table_text
        assert "RBTL" in table_text


def test_end_sem_download_places_or_between_alternative_questions() -> None:
    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        generate_response = client.post(
            "/api/v1/ai/generate-paper",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": 3,
                "title": "End-Sem NLP Paper",
                "exam_type": "End-Sem",
                "semester": "6",
                "batch": "2022-26",
                "max_marks": 100,
                "duration_minutes": 180,
                "exam_date": "2026-04-30",
                "teaching_department": "AIML",
                "prompt": "Generate a module-balanced end-sem paper for NLP.",
                "rbt_levels": ["L1", "L2", "L3", "L4", "L5", "L6"],
                "module_numbers": [1, 2, 3, 4, 5],
                "difficulty_distribution": {"easy": 30, "medium": 40, "hard": 30},
                "co_targets": {"CO1": 20, "CO2": 20, "CO3": 20, "CO4": 20, "CO5": 20},
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()
        assert len(paper["questions"]) == 26

        download_response = client.get(
            f"/api/v1/papers/{paper['id']}/download",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert download_response.status_code == 200
        exported = DocxDocument(BytesIO(download_response.content))
        question_table = _question_table(exported)

        labels = [row.cells[0].text.strip() for row in question_table.rows]
        question_rows = {
            row.cells[0].text.strip(): row.cells[1].text.strip()
            for row in question_table.rows
            if row.cells[0].text.strip() and row.cells[0].text.strip()[0].isdigit()
        }
        assert len(question_rows) == 26
        assert all(question_rows.values())

        expected_breaks = [
            ("1c", "2a"),
            ("3c", "4a"),
            ("5c", "6a"),
            ("7b", "8a"),
            ("9b", "10a"),
        ]
        for previous_label, next_label in expected_breaks:
            previous_index = labels.index(previous_label)
            assert labels[previous_index + 1] == "OR"
            assert labels[previous_index + 2] == next_label


def test_manual_generation_respects_selected_question_order() -> None:
    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        questions_response = client.get(
            "/api/v1/questions?subject_id=1",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert questions_response.status_code == 200
        manual_ids = [item["id"] for item in questions_response.json()[:20]]
        assert len(manual_ids) == 20

        generate_response = client.post(
            "/api/v1/ai/generate-paper",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": 1,
                "title": "Manual ML Paper",
                "exam_type": "IAT-2",
                "semester": "5",
                "batch": "2022-26",
                "max_marks": 50,
                "duration_minutes": 90,
                "teaching_department": "AIML",
                "prompt": "Use the manually curated question list.",
                "rbt_levels": ["L1", "L2", "L3", "L4"],
                "module_numbers": [1, 2, 3, 4, 5],
                "manual_question_ids": manual_ids,
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()
        assert [item["question_id"] for item in paper["questions"]] == manual_ids
