from __future__ import annotations

import base64
import io
import json
import re
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

import httpx
from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader
from sqlalchemy import func, select
from sqlalchemy.orm import Session, selectinload

from .auth import create_auth_tokens, hash_password, verify_password
from .config import settings
from .generator import PaperConfig, generate_question_paper
from .models import (
    AuditLog,
    Department,
    Document,
    PaperQuestion,
    PaperReview,
    PaperStatus,
    Question,
    QuestionPaper,
    ReviewDecision,
    Role,
    Subject,
    TeacherSubject,
    User,
)


def authenticate_user(db: Session, email: str, password: str) -> dict[str, str]:
    user = db.scalar(select(User).where(User.email == email))
    if user is None or not verify_password(password, user.password_hash):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid credentials"
        )
    user.last_login_at = datetime.now(timezone.utc)
    db.commit()
    return create_auth_tokens(user)


def write_audit_log(
    db: Session,
    user_id: int | None,
    action: str,
    entity: str,
    entity_id: str | None,
    details: dict[str, Any],
) -> None:
    db.add(
        AuditLog(
            user_id=user_id,
            action=action,
            entity=entity,
            entity_id=entity_id,
            details=details,
        )
    )
    db.commit()


def ensure_subject_access(user: User, subject: Subject, db: Session) -> None:
    if user.role == Role.ADMIN:
        return
    if user.dept_id != subject.dept_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN, detail="Department access denied"
        )
    if user.role == Role.TEACHER:
        link = db.scalar(
            select(TeacherSubject).where(
                TeacherSubject.teacher_id == user.id,
                TeacherSubject.subject_id == subject.id,
            )
        )
        if link is None:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Teacher is not assigned to this subject",
            )


import logging

logger = logging.getLogger("app.services")


class OllamaService:
    def __init__(self) -> None:
        self.base_url = settings.ollama_base_url.rstrip("/")
        self.model = settings.ollama_model
        logger.info(
            f"OllamaService initialized with base_url={self.base_url}, model={self.model}"
        )

    def _call(
        self, prompt: str, system: str, images: list[str] | None = None
    ) -> dict[str, Any] | None:
        payload: dict[str, Any] = {
            "model": self.model,
            "prompt": prompt,
            "system": system,
            "stream": False,
            "format": "json",
        }
        if images:
            payload["images"] = images

        try:
            logger.debug(f"Calling Ollama API at {self.base_url}/api/generate")
            response = httpx.post(
                f"{self.base_url}/api/generate", json=payload, timeout=45.0
            )
            response.raise_for_status()
            raw = response.json().get("response", "{}")
            logger.debug(f"Ollama response received: {raw[:100]}...")
            return json.loads(raw)
        except Exception as e:
            logger.error(f"Ollama API call failed: {str(e)}")
            return None

    def extract_questions_from_text(self, text: str) -> list[dict[str, Any]] | None:
        logger.info("Extracting questions from text using AI")
        prompt = (
            "Extract academic questions from the input. Return JSON with a `questions` array. "
            "Each item must include text, marks, bloom_level, difficulty, module_number, course_outcome, and tags."
            f"\n\nINPUT:\n{text[:12000]}"
        )
        system = (
            "You are an academic parser for engineering question banks. "
            "Return only valid JSON. Use bloom levels L1-L6 and concise tags."
        )
        data = self._call(prompt, system)
        if isinstance(data, dict) and isinstance(data.get("questions"), list):
            logger.info(f"Successfully extracted {len(data['questions'])} questions")
            return data["questions"]
        logger.warning("AI extraction failed or returned invalid format")
        return None

    def extract_questions_from_image(
        self, content: bytes
    ) -> list[dict[str, Any]] | None:
        logger.info("Extracting questions from image using AI Vision")
        encoded = base64.b64encode(content).decode("utf-8")
        prompt = (
            "Read the image and extract academic question bank items. "
            "Return JSON with a `questions` array using the same schema as text extraction."
        )
        system = "You are using LLaMA 3.2 Vision to extract exam questions from uploaded source material."
        data = self._call(prompt, system, images=[encoded])
        if isinstance(data, dict) and isinstance(data.get("questions"), list):
            logger.info(
                f"Successfully extracted {len(data['questions'])} questions from image"
            )
            return data["questions"]
        logger.warning("AI vision extraction failed")
        return None

    def generate_summary(
        self, prompt: str, questions: list[Question], max_marks: int
    ) -> str | None:
        logger.info("Generating paper summary using AI")
        question_lines = "\n".join(
            f"- [{question.id}] {question.text} ({question.marks} marks, {question.course_outcome}, {question.bloom_level})"
            for question in questions
        )
        user_prompt = (
            "Summarize the generated paper in 2-3 sentences for a teacher preview. "
            f"Prompt: {prompt}\nTarget marks: {max_marks}\nSelected questions:\n{question_lines}"
        )
        system = "You are an assistant explaining how a question paper matches academic constraints."
        data = self._call(user_prompt, system)
        if isinstance(data, dict) and isinstance(data.get("summary"), str):
            return data["summary"]
        logger.warning("AI summary generation failed")
        return None


ollama_service = OllamaService()

RBT_DESCRIPTIONS = {
    "L1": "Remember",
    "L2": "Understand",
    "L3": "Apply",
    "L4": "Analyze",
    "L5": "Evaluate",
    "L6": "Create",
}


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    font_size: int = 10,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_doc_defaults(doc: DocxDocument) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.start_type = WD_SECTION.NEW_PAGE

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _add_paragraph(
    doc: DocxDocument,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    font_size: int = 10,
    space_after: int = 0,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _set_table_borders(table) -> None:
    table_element = table._tbl
    table_properties = table_element.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _question_suffix(index: int, total: int) -> str:
    if index % 2 == 1 and index < total:
        return "\nOR"
    return ""


def _format_duration(duration_minutes: int) -> str:
    return f"{duration_minutes} Minutes"


def _guess_bloom_level(text: str) -> str:
    lowered = text.lower()
    mapping = {
        "define": "L1",
        "list": "L1",
        "explain": "L2",
        "describe": "L2",
        "illustrate": "L3",
        "apply": "L3",
        "analyze": "L4",
        "compare": "L4",
        "evaluate": "L5",
        "justify": "L5",
        "design": "L6",
        "develop": "L6",
    }
    for key, value in mapping.items():
        if key in lowered:
            return value
    return "L2"


def _guess_difficulty(marks: int, bloom_level: str) -> str:
    if marks <= 5 and bloom_level in {"L1", "L2"}:
        return "easy"
    if marks >= 10 or bloom_level in {"L5", "L6"}:
        return "hard"
    return "medium"


def _extract_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        return content.decode("utf-8", errors="ignore")
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".docx":
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if suffix in {".png", ".jpg", ".jpeg"}:
        return ""
    return content.decode("utf-8", errors="ignore")


def _fallback_question_parse(text: str) -> list[dict[str, Any]]:
    candidates = [
        line.strip(" -\t") for line in re.split(r"\n+", text) if len(line.strip()) > 15
    ]
    questions: list[dict[str, Any]] = []
    for line in candidates:
        clean = re.sub(r"^\d+[\).\s-]*", "", line).strip()
        marks = 10 if len(clean) > 100 else 5
        bloom = _guess_bloom_level(clean)
        questions.append(
            {
                "text": clean,
                "marks": marks,
                "bloom_level": bloom,
                "difficulty": _guess_difficulty(marks, bloom),
                "module_number": 1,
                "course_outcome": "CO1",
                "tags": ["uploaded"],
            }
        )
    return questions[:50]


def parse_uploaded_document(
    db: Session, user: User, subject_id: int, upload: UploadFile
) -> tuple[Document, list[Question], str]:
    subject = db.get(Subject, subject_id)
    if subject is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Subject not found"
        )
    ensure_subject_access(user, subject, db)

    content = upload.file.read()
    subject_dir = settings.storage_path / "uploads" / str(subject_id)
    subject_dir.mkdir(parents=True, exist_ok=True)
    saved_path = subject_dir / f"{uuid4().hex}_{upload.filename}"
    saved_path.write_bytes(content)

    suffix = Path(upload.filename or "").suffix.lower()
    parsed_text = _extract_text(upload.filename or "upload.txt", content)

    extracted = None
    ai_mode = "heuristic"
    if suffix in {".png", ".jpg", ".jpeg"}:
        extracted = ollama_service.extract_questions_from_image(content)
        parsed_text = parsed_text or f"Image upload: {upload.filename}"
        ai_mode = "vision" if extracted else "heuristic"
    elif parsed_text.strip():
        extracted = ollama_service.extract_questions_from_text(parsed_text)
        ai_mode = "llm-text" if extracted else "heuristic"

    parsed_questions = extracted or _fallback_question_parse(
        parsed_text or upload.filename or "Untitled upload"
    )
    if not parsed_questions:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not extract questions from file",
        )

    document = Document(
        subject_id=subject.id,
        teacher_id=user.id,
        filename=upload.filename or "upload",
        mime_type=upload.content_type or "application/octet-stream",
        storage_path=str(saved_path),
        parsed_text=parsed_text[:50000],
        upload_status="processed",
    )
    db.add(document)
    db.flush()

    question_rows: list[Question] = []
    for item in parsed_questions:
        text = str(item.get("text", "")).strip()
        if not text:
            continue
        marks = int(item.get("marks", 5))
        bloom = str(item.get("bloom_level", _guess_bloom_level(text))).upper()
        question_rows.append(
            Question(
                subject_id=subject.id,
                teacher_id=user.id,
                source_doc_id=document.id,
                text=text,
                marks=marks,
                course_outcome=str(item.get("course_outcome", "CO1")).upper(),
                bloom_level=bloom,
                difficulty=str(
                    item.get("difficulty", _guess_difficulty(marks, bloom))
                ).lower(),
                module_number=int(item.get("module_number", 1)),
                tags=[str(tag) for tag in item.get("tags", ["uploaded"])],
                is_verified=False,
            )
        )

    db.add_all(question_rows)
    db.commit()
    write_audit_log(
        db,
        user.id,
        "upload_questions",
        "document",
        str(document.id),
        {"filename": document.filename, "count": len(question_rows)},
    )
    return document, question_rows, ai_mode


def create_question(db: Session, user: User, payload: dict[str, Any]) -> Question:
    subject = db.get(Subject, payload["subject_id"])
    if subject is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Subject not found"
        )
    ensure_subject_access(user, subject, db)
    question = Question(
        teacher_id=user.id, is_verified=user.role != Role.TEACHER, **payload
    )
    db.add(question)
    db.commit()
    db.refresh(question)
    write_audit_log(
        db,
        user.id,
        "create_question",
        "question",
        str(question.id),
        {"subject_id": subject.id},
    )
    return question


def update_question(
    db: Session, user: User, question_id: int, payload: dict[str, Any]
) -> Question:
    question = db.get(Question, question_id)
    if question is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Question not found"
        )
    if user.role == Role.TEACHER and question.teacher_id != user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only the owning teacher can edit this question",
        )
    payload["teacher_id"] = question.teacher_id
    for key, value in payload.items():
        setattr(question, key, value)
    db.commit()
    db.refresh(question)
    write_audit_log(
        db,
        user.id,
        "update_question",
        "question",
        str(question.id),
        {"subject_id": question.subject_id},
    )
    return question


def delete_question(db: Session, user: User, question_id: int) -> None:
    question = db.get(Question, question_id)
    if question is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Question not found"
        )
    if user.role == Role.TEACHER and question.teacher_id != user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only the owning teacher can delete this question",
        )
    db.delete(question)
    db.commit()
    write_audit_log(
        db,
        user.id,
        "delete_question",
        "question",
        str(question_id),
        {"subject_id": question.subject_id},
    )


def list_questions_for_user(
    db: Session,
    user: User,
    search: str | None,
    subject_id: int | None,
    bloom_level: str | None,
    difficulty: str | None,
) -> list[Question]:
    stmt = select(Question)
    if user.role == Role.TEACHER:
        stmt = stmt.where(Question.teacher_id == user.id)
    elif user.role == Role.HOD and user.dept_id is not None:
        subject_ids = select(Subject.id).where(Subject.dept_id == user.dept_id)
        stmt = stmt.where(Question.subject_id.in_(subject_ids))

    if subject_id:
        stmt = stmt.where(Question.subject_id == subject_id)
    if search:
        stmt = stmt.where(Question.text.ilike(f"%{search}%"))
    if bloom_level:
        stmt = stmt.where(Question.bloom_level == bloom_level.upper())
    if difficulty:
        stmt = stmt.where(Question.difficulty == difficulty.lower())
    return list(db.scalars(stmt.order_by(Question.created_at.desc())))


def _select_questions(
    candidates: list[Question],
    max_marks: int,
    requested_bloom: set[str],
    modules: set[int],
) -> list[Question]:
    filtered = [
        question
        for question in candidates
        if (not requested_bloom or question.bloom_level in requested_bloom)
        and (not modules or question.module_number in modules)
    ]
    if not filtered:
        filtered = candidates

    grouped: dict[tuple[str, int], list[Question]] = defaultdict(list)
    for question in filtered:
        grouped[(question.course_outcome, question.module_number)].append(question)

    for bucket in grouped.values():
        bucket.sort(key=lambda item: (item.marks, item.created_at), reverse=True)

    selected: list[Question] = []
    total = 0
    seen_ids: set[int] = set()
    for _, bucket in sorted(
        grouped.items(), key=lambda entry: (entry[0][1], entry[0][0])
    ):
        for question in bucket:
            if question.id in seen_ids:
                continue
            if total + question.marks > max_marks:
                continue
            selected.append(question)
            total += question.marks
            seen_ids.add(question.id)
            break
        if total >= max_marks:
            return selected

    remaining = sorted(
        [question for question in filtered if question.id not in seen_ids],
        key=lambda item: item.marks,
    )
    while total < max_marks and remaining:
        added = False
        for question in list(remaining):
            if total + question.marks <= max_marks:
                selected.append(question)
                total += question.marks
                remaining.remove(question)
                added = True
                break
        if not added:
            break

    return selected


def generate_paper(db: Session, user: User, payload: dict[str, Any]) -> QuestionPaper:
    subject = db.get(Subject, payload["subject_id"])
    if subject is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Subject not found"
        )
    ensure_subject_access(user, subject, db)

    candidates = list(
        db.scalars(
            select(Question)
            .where(Question.subject_id == subject.id)
            .order_by(Question.module_number.asc(), Question.created_at.asc())
        )
    )
    if not candidates:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No questions available for this subject",
        )

    selected = _select_questions(
        candidates,
        payload["max_marks"],
        {level.upper() for level in payload.get("rbt_levels", [])},
        set(payload.get("module_numbers", [])),
    )
    if not selected:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Could not assemble a paper from the available questions",
        )

    summary = ollama_service.generate_summary(
        payload["prompt"], selected, payload["max_marks"]
    ) or (
        f"Generated {len(selected)} questions for {subject.code} using the local selection engine. "
        f"Coverage spans {len({question.course_outcome for question in selected})} course outcomes."
    )

    coverage_stats = {
        "by_module": dict(
            sorted(Counter(question.module_number for question in selected).items())
        ),
        "by_rbt": dict(
            sorted(Counter(question.bloom_level for question in selected).items())
        ),
        "by_co": dict(
            sorted(Counter(question.course_outcome for question in selected).items())
        ),
    }

    paper = QuestionPaper(
        subject_id=subject.id,
        teacher_id=user.id,
        title=payload["title"],
        exam_type=payload["exam_type"],
        semester=payload["semester"],
        batch=payload["batch"],
        max_marks=payload["max_marks"],
        duration_minutes=payload["duration_minutes"],
        exam_date=payload.get("exam_date"),
        teaching_department=payload["teaching_department"],
        prompt_used=payload["prompt"],
        ai_config_json={
            "model": settings.ollama_model,
            "rbt_levels": payload.get("rbt_levels", []),
            "module_numbers": payload.get("module_numbers", []),
            "difficulty_distribution": payload.get("difficulty_distribution", {}),
            "co_targets": payload.get("co_targets", {}),
            "co_descriptions": payload.get("co_descriptions", {}),
            "difficulty": payload.get("difficulty", "balanced"),
            "instructions": payload.get(
                "instructions", "Instruction: Answer the following questions"
            ),
            "coverage_stats": coverage_stats,
            "allow_ai_rewrite": payload.get("allow_ai_rewrite", False),
        },
        generated_summary=summary,
    )
    db.add(paper)
    db.flush()

    for index, question in enumerate(selected, start=1):
        db.add(
            PaperQuestion(
                paper_id=paper.id,
                question_id=question.id,
                order_index=index,
                section_label="A",
                custom_marks=question.marks,
                question_text_snapshot=question.text,
            )
        )

    db.commit()
    db.refresh(paper)
    write_audit_log(
        db,
        user.id,
        "generate_paper",
        "paper",
        str(paper.id),
        {"subject_id": subject.id, "question_count": len(selected)},
    )
    return paper


def _paper_context(
    db: Session, paper: QuestionPaper
) -> tuple[
    Subject | None,
    Department | None,
    dict[int, Question],
    list[dict[str, Any]],
    dict[str, Any],
]:
    ordered = sorted(paper.questions, key=lambda item: item.order_index)
    subject = db.get(Subject, paper.subject_id)
    department = db.get(Department, subject.dept_id) if subject else None
    question_ids = [item.question_id for item in ordered]
    question_rows = (
        list(db.scalars(select(Question).where(Question.id.in_(question_ids))))
        if question_ids
        else []
    )
    questions_by_id = {question.id: question for question in question_rows}
    serialized_questions = [
        {
            "id": item.id,
            "question_id": item.question_id,
            "order_index": item.order_index,
            "section_label": item.section_label,
            "custom_marks": item.custom_marks,
            "text": item.question_text_snapshot,
            "course_outcome": questions_by_id.get(item.question_id).course_outcome
            if questions_by_id.get(item.question_id)
            else None,
            "bloom_level": questions_by_id.get(item.question_id).bloom_level
            if questions_by_id.get(item.question_id)
            else None,
            "module_number": questions_by_id.get(item.question_id).module_number
            if questions_by_id.get(item.question_id)
            else None,
            "difficulty": questions_by_id.get(item.question_id).difficulty
            if questions_by_id.get(item.question_id)
            else None,
        }
        for item in ordered
    ]
    ai_config = paper.ai_config_json or {}
    coverage_stats = ai_config.get("coverage_stats", {})
    return subject, department, questions_by_id, serialized_questions, coverage_stats


def serialize_paper(db: Session, paper: QuestionPaper) -> dict[str, Any]:
    subject, department, _, serialized_questions, coverage_stats = _paper_context(
        db, paper
    )
    return {
        "id": paper.id,
        "subject_id": paper.subject_id,
        "subject_name": subject.name if subject else None,
        "subject_code": subject.code if subject else None,
        "department_name": department.name if department else None,
        "title": paper.title,
        "exam_type": paper.exam_type,
        "semester": paper.semester,
        "batch": paper.batch,
        "max_marks": paper.max_marks,
        "duration_minutes": paper.duration_minutes,
        "exam_date": paper.exam_date,
        "teaching_department": paper.teaching_department,
        "status": paper.status,
        "prompt_used": paper.prompt_used,
        "generated_summary": paper.generated_summary,
        "created_at": paper.created_at,
        "submitted_at": paper.submitted_at,
        "reviewed_at": paper.reviewed_at,
        "ai_config": paper.ai_config_json or {},
        "coverage_stats": coverage_stats,
        "questions": serialized_questions,
    }


def get_paper_or_404(db: Session, paper_id: int) -> QuestionPaper:
    paper = db.scalar(select(QuestionPaper).where(QuestionPaper.id == paper_id))
    if paper is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found"
        )
    return paper


def ensure_paper_access(db: Session, user: User, paper: QuestionPaper) -> None:
    subject = db.get(Subject, paper.subject_id)
    if subject is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Paper subject missing"
        )
    if user.role == Role.ADMIN:
        return
    if user.role == Role.TEACHER and paper.teacher_id != user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN, detail="Paper ownership required"
        )
    if user.role == Role.HOD and user.dept_id != subject.dept_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN, detail="Department access denied"
        )


def list_papers_for_user(db: Session, user: User) -> list[QuestionPaper]:
    stmt = select(QuestionPaper).options(selectinload(QuestionPaper.questions))
    if user.role == Role.TEACHER:
        stmt = stmt.where(QuestionPaper.teacher_id == user.id)
    elif user.role == Role.HOD and user.dept_id is not None:
        subject_ids = select(Subject.id).where(Subject.dept_id == user.dept_id)
        stmt = stmt.where(QuestionPaper.subject_id.in_(subject_ids))
    return list(db.scalars(stmt.order_by(QuestionPaper.created_at.desc())))


def update_paper(
    db: Session,
    user: User,
    paper: QuestionPaper,
    title: str | None,
    prompt: str | None,
    overrides: dict[int, str],
) -> QuestionPaper:
    ensure_paper_access(db, user, paper)
    if user.role != Role.TEACHER or paper.teacher_id != user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only the owning teacher can edit this paper",
        )
    if paper.status not in {PaperStatus.DRAFT, PaperStatus.REJECTED}:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Paper can no longer be edited",
        )

    if title:
        paper.title = title
    if prompt:
        paper.prompt_used = prompt
    if overrides:
        for item in paper.questions:
            if item.id in overrides:
                item.question_text_snapshot = overrides[item.id]
    db.commit()
    db.refresh(paper)
    write_audit_log(
        db,
        user.id,
        "update_paper",
        "paper",
        str(paper.id),
        {"override_count": len(overrides)},
    )
    return paper


def submit_paper(db: Session, user: User, paper: QuestionPaper) -> QuestionPaper:
    ensure_paper_access(db, user, paper)
    if user.role != Role.TEACHER or paper.teacher_id != user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only the owning teacher can submit this paper",
        )
    paper.status = PaperStatus.PENDING_REVIEW
    paper.submitted_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(paper)
    write_audit_log(db, user.id, "submit_paper", "paper", str(paper.id), {})
    return paper


def review_paper(
    db: Session,
    user: User,
    paper: QuestionPaper,
    decision: ReviewDecision,
    comments: str,
) -> QuestionPaper:
    ensure_paper_access(db, user, paper)
    if user.role not in {Role.HOD, Role.ADMIN}:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN, detail="Review permission required"
        )
    if paper.status != PaperStatus.PENDING_REVIEW:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Paper is not pending review",
        )

    paper.status = (
        PaperStatus.APPROVED
        if decision == ReviewDecision.APPROVED
        else PaperStatus.REJECTED
    )
    paper.reviewed_at = datetime.now(timezone.utc)
    db.add(
        PaperReview(
            paper_id=paper.id, reviewer_id=user.id, status=decision, comments=comments
        )
    )
    db.commit()
    db.refresh(paper)
    write_audit_log(
        db,
        user.id,
        "review_paper",
        "paper",
        str(paper.id),
        {"decision": decision, "comments": comments},
    )
    return paper


def export_paper_docx(db: Session, user: User, paper: QuestionPaper) -> Path:
    ensure_paper_access(db, user, paper)
    export_dir = settings.storage_path / "papers"
    export_dir.mkdir(parents=True, exist_ok=True)
    subject, department, _, serialized_questions, coverage_stats = _paper_context(
        db, paper
    )
    ai_config = paper.ai_config_json or {}
    co_percentages = (
        coverage_stats.get("percentages", {}).get("co", {})
        if coverage_stats
        else {}
    )
    module_percentages = (
        coverage_stats.get("percentages", {}).get("modules", {})
        if coverage_stats
        else {}
    )

    config = PaperConfig(
        department=department.name if department else paper.teaching_department,
        subject=subject.name if subject else "Subject",
        subject_code=subject.code if subject else "N/A",
        semester=paper.semester,
        max_marks=paper.max_marks,
        duration=_format_duration(paper.duration_minutes),
        date=paper.exam_date.strftime("%d-%m-%Y") if paper.exam_date else "TBD",
        batch=paper.batch,
        teaching_department=paper.teaching_department,
        exam_type=paper.exam_type,
        modules=[int(value) for value in ai_config.get("module_numbers", []) if str(value).strip()],
        rbt_levels=[str(value).upper() for value in ai_config.get("rbt_levels", []) if str(value).strip()],
        co_targets=[str(value).upper() for value in ai_config.get("co_targets", {}).keys()],
        instructions=ai_config.get(
            "instructions", "Instruction: Answer the following questions"
        ),
        co_descriptions={
            str(key).upper(): str(value)
            for key, value in ai_config.get("co_descriptions", {}).items()
        },
        co_percentages={str(key).upper(): int(value) for key, value in co_percentages.items()},
        module_percentages={str(key): int(value) for key, value in module_percentages.items()},
    )

    marks_per_slot = max(1, paper.max_marks // 5) if paper.max_marks else 10
    docx_questions = [
        {
            "text": question["text"],
            "marks": question["custom_marks"] or marks_per_slot,
            "course_outcome": question.get("course_outcome") or "",
            "bloom_level": question.get("bloom_level") or "",
            "module_number": question.get("module_number"),
        }
        for question in serialized_questions
    ]

    file_path = generate_question_paper(config, docx_questions, export_dir)
    paper.download_path = str(file_path)
    if paper.status == PaperStatus.APPROVED:
        paper.status = PaperStatus.PRINT_READY
    db.commit()
    write_audit_log(
        db, user.id, "export_paper", "paper", str(paper.id), {"path": str(file_path)}
    )
    return file_path


def delete_paper(db: Session, user: User, paper: QuestionPaper) -> None:
    ensure_paper_access(db, user, paper)
    if user.role == Role.TEACHER and paper.teacher_id != user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only the owning teacher can delete this paper",
        )
    db.delete(paper)
    db.commit()
    write_audit_log(db, user.id, "delete_paper", "paper", str(paper.id), {})


def create_admin_user(db: Session, actor: User, payload: dict[str, Any]) -> User:
    if db.scalar(select(User).where(User.email == payload["email"])) is not None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="User already exists"
        )
    user = User(
        email=payload["email"],
        full_name=payload["full_name"],
        password_hash=hash_password(payload["password"]),
        role=payload["role"],
        dept_id=payload.get("dept_id"),
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    write_audit_log(
        db, actor.id, "create_user", "user", str(user.id), {"role": user.role}
    )
    return user


def dashboard_stats(db: Session) -> dict[str, Any]:
    return {
        "total_users": db.scalar(select(func.count()).select_from(User)) or 0,
        "total_subjects": db.scalar(select(func.count()).select_from(Subject)) or 0,
        "total_questions": db.scalar(select(func.count()).select_from(Question)) or 0,
        "total_papers": db.scalar(select(func.count()).select_from(QuestionPaper)) or 0,
        "pending_reviews": db.scalar(
            select(func.count())
            .select_from(QuestionPaper)
            .where(QuestionPaper.status == PaperStatus.PENDING_REVIEW)
        )
        or 0,
        "approved_papers": db.scalar(
            select(func.count())
            .select_from(QuestionPaper)
            .where(
                QuestionPaper.status.in_(
                    [PaperStatus.APPROVED, PaperStatus.PRINT_READY]
                )
            )
        )
        or 0,
        "ai_model": settings.ollama_model,
    }


def seed_demo_data(db: Session) -> None:
    if not settings.allow_demo_seed:
        return

    def ensure_user(
        email: str,
        full_name: str,
        role: Role,
        dept_id: int,
        password: str,
    ) -> User:
        user = db.scalar(select(User).where(User.email == email))
        if user is None:
            user = User(
                email=email,
                full_name=full_name,
                password_hash=hash_password(password),
                role=role,
                dept_id=dept_id,
            )
            db.add(user)
            db.flush()
        return user

    def build_module_questions(
        subject_name: str,
        module_title: str,
        focus_terms: list[str],
        module_number: int,
    ) -> list[tuple[str, int, str, str, str, int]]:
        focus_a, focus_b, focus_c, focus_d = focus_terms
        subject_area = subject_name.lower()
        return [
            (
                f"Explain {focus_a} and its role in {module_title.lower()}.",
                5,
                "CO1",
                "L2",
                "easy",
                module_number,
            ),
            (
                f"Compare {focus_a} and {focus_b} for {module_title.lower()} tasks.",
                5,
                "CO2",
                "L4",
                "medium",
                module_number,
            ),
            (
                f"Illustrate the workflow of {focus_c} with a neat diagram and suitable example.",
                10,
                "CO2",
                "L3",
                "medium",
                module_number,
            ),
            (
                f"Apply {focus_b} to solve a representative problem in {subject_area}.",
                10,
                "CO3",
                "L3",
                "medium",
                module_number,
            ),
            (
                f"Analyze design trade-offs in {focus_d} for scalable {subject_area} systems.",
                10,
                "CO4",
                "L4",
                "hard",
                module_number,
            ),
            (
                f"Design an end-to-end {subject_area} solution using {focus_a}, {focus_b}, and {focus_c}.",
                10,
                "CO5",
                "L6",
                "hard",
                module_number,
            ),
            (
                f"Justify how {focus_d} improves reliability or accuracy in {module_title.lower()}.",
                5,
                "CO3",
                "L5",
                "medium",
                module_number,
            ),
            (
                f"Evaluate the limitations of {focus_c} and propose suitable improvements.",
                10,
                "CO5",
                "L5",
                "hard",
                module_number,
            ),
        ]

    dept = db.scalar(select(Department).where(Department.code == "AIML"))
    if dept is None:
        dept = Department(
            name="Artificial Intelligence & Machine Learning",
            code="AIML",
        )
        db.add(dept)
        db.flush()

    teacher = ensure_user(
        "teacher@dsatm.edu",
        "Dr. Asha Rao",
        Role.TEACHER,
        dept.id,
        "Teacher@123",
    )
    ensure_user("hod@dsatm.edu", "Prof. Vinay Kumar", Role.HOD, dept.id, "Hod@123")
    ensure_user(
        "admin@dsatm.edu",
        "System Administrator",
        Role.ADMIN,
        dept.id,
        "Admin@123",
    )

    subject_catalog = [
        {
            "name": "Machine Learning",
            "code": "21AI51",
            "semester": 5,
            "modules": [
                (
                    "Foundations of Machine Learning",
                    [
                        "supervised learning",
                        "unsupervised learning",
                        "bias-variance trade-off",
                        "feature scaling",
                    ],
                ),
                (
                    "Regression and Optimization",
                    [
                        "linear regression",
                        "gradient descent",
                        "regularization",
                        "logistic regression",
                    ],
                ),
                (
                    "Classification and Ensembles",
                    [
                        "decision trees",
                        "support vector machines",
                        "naive Bayes classifiers",
                        "random forests",
                    ],
                ),
                (
                    "Clustering and Dimensionality Reduction",
                    [
                        "k-means clustering",
                        "DBSCAN",
                        "principal component analysis",
                        "feature extraction",
                    ],
                ),
                (
                    "Evaluation and Deployment",
                    [
                        "cross-validation",
                        "ROC-AUC analysis",
                        "model deployment pipelines",
                        "hyperparameter tuning",
                    ],
                ),
            ],
        },
        {
            "name": "Deep Learning",
            "code": "21AI52",
            "semester": 5,
            "modules": [
                (
                    "Neural Network Basics",
                    [
                        "perceptron models",
                        "activation functions",
                        "backpropagation",
                        "loss landscapes",
                    ],
                ),
                (
                    "Convolutional Neural Networks",
                    [
                        "convolution layers",
                        "pooling strategies",
                        "transfer learning",
                        "image classification pipelines",
                    ],
                ),
                (
                    "Sequence Models",
                    [
                        "recurrent neural networks",
                        "LSTM cells",
                        "GRU networks",
                        "attention mechanisms",
                    ],
                ),
                (
                    "Generative Deep Learning",
                    [
                        "GAN training",
                        "variational autoencoders",
                        "diffusion models",
                        "style transfer",
                    ],
                ),
                (
                    "Advanced Deep Learning Topics",
                    [
                        "few-shot learning",
                        "model compression",
                        "multimodal learning",
                        "ethical AI deployment",
                    ],
                ),
            ],
        },
        {
            "name": "Natural Language Processing",
            "code": "21AI53",
            "semester": 6,
            "modules": [
                (
                    "Text Processing",
                    [
                        "tokenization",
                        "stemming and lemmatization",
                        "part-of-speech tagging",
                        "named entity recognition",
                    ],
                ),
                (
                    "Word Representations",
                    [
                        "n-gram language models",
                        "Word2Vec embeddings",
                        "GloVe embeddings",
                        "fastText representations",
                    ],
                ),
                (
                    "Sequence Modelling",
                    [
                        "encoder-decoder architectures",
                        "beam search decoding",
                        "Bi-LSTM tagging",
                        "attention-based sequence modelling",
                    ],
                ),
                (
                    "Transformers",
                    [
                        "transformer encoders",
                        "self-attention",
                        "BERT fine-tuning",
                        "GPT style generation",
                    ],
                ),
                (
                    "Advanced NLP Applications",
                    [
                        "machine translation",
                        "question answering systems",
                        "text summarization",
                        "cross-lingual transfer",
                    ],
                ),
            ],
        },
        {
            "name": "Computer Vision",
            "code": "21AI54",
            "semester": 6,
            "modules": [
                (
                    "Image Foundations",
                    [
                        "image acquisition",
                        "image enhancement",
                        "histogram equalization",
                        "edge detection",
                    ],
                ),
                (
                    "Feature Extraction",
                    [
                        "SIFT descriptors",
                        "HOG features",
                        "optical flow",
                        "image segmentation",
                    ],
                ),
                (
                    "Detection and Recognition",
                    [
                        "object detection",
                        "region proposal networks",
                        "face recognition",
                        "metric learning",
                    ],
                ),
                (
                    "Deep Vision Models",
                    [
                        "vision transformers",
                        "semantic segmentation",
                        "instance segmentation",
                        "transfer learning for vision",
                    ],
                ),
                (
                    "Applied Computer Vision",
                    [
                        "medical image analysis",
                        "autonomous driving perception",
                        "video analytics",
                        "edge deployment",
                    ],
                ),
            ],
        },
        {
            "name": "Data Mining and Data Warehousing",
            "code": "21AI61",
            "semester": 6,
            "modules": [
                (
                    "Data Warehousing",
                    [
                        "ETL pipelines",
                        "fact and dimension schemas",
                        "OLAP cubes",
                        "data cleaning",
                    ],
                ),
                (
                    "Association and Classification",
                    [
                        "association rule mining",
                        "Apriori algorithm",
                        "decision-tree classification",
                        "Bayesian classification",
                    ],
                ),
                (
                    "Clustering and Outlier Analysis",
                    [
                        "hierarchical clustering",
                        "density-based clustering",
                        "outlier detection",
                        "cluster validation",
                    ],
                ),
                (
                    "Pattern Evaluation",
                    [
                        "interestingness measures",
                        "time-series mining",
                        "web usage mining",
                        "text mining",
                    ],
                ),
                (
                    "Analytics Applications",
                    [
                        "business intelligence dashboards",
                        "stream mining",
                        "recommendation systems",
                        "privacy-preserving mining",
                    ],
                ),
            ],
        },
        {
            "name": "Reinforcement Learning",
            "code": "21AI62",
            "semester": 6,
            "modules": [
                (
                    "RL Foundations",
                    [
                        "Markov decision processes",
                        "reward design",
                        "value functions",
                        "policy evaluation",
                    ],
                ),
                (
                    "Dynamic Programming and Monte Carlo Methods",
                    [
                        "Bellman equations",
                        "policy iteration",
                        "Monte Carlo control",
                        "temporal-difference learning",
                    ],
                ),
                (
                    "Model-Free Control",
                    [
                        "Q-learning",
                        "SARSA",
                        "exploration strategies",
                        "function approximation",
                    ],
                ),
                (
                    "Deep Reinforcement Learning",
                    [
                        "deep Q-networks",
                        "policy gradient methods",
                        "actor-critic algorithms",
                        "experience replay",
                    ],
                ),
                (
                    "Advanced Reinforcement Learning",
                    [
                        "multi-agent reinforcement learning",
                        "hierarchical RL",
                        "safe reinforcement learning",
                        "offline RL",
                    ],
                ),
            ],
        },
        {
            "name": "Generative AI",
            "code": "21AI71",
            "semester": 7,
            "modules": [
                (
                    "Language Model Foundations",
                    [
                        "token prediction",
                        "transformer pretraining",
                        "prompt engineering",
                        "embedding search",
                    ],
                ),
                (
                    "LLM Adaptation",
                    [
                        "instruction tuning",
                        "retrieval-augmented generation",
                        "parameter-efficient fine-tuning",
                        "evaluation benchmarks",
                    ],
                ),
                (
                    "Multimodal Generation",
                    [
                        "vision-language models",
                        "image captioning",
                        "text-to-image generation",
                        "speech interfaces",
                    ],
                ),
                (
                    "Responsible Generative AI",
                    [
                        "hallucination mitigation",
                        "safety alignment",
                        "copyright and provenance",
                        "human feedback loops",
                    ],
                ),
                (
                    "Generative AI Deployment",
                    [
                        "agent orchestration",
                        "latency optimization",
                        "guardrails",
                        "monitoring and observability",
                    ],
                ),
            ],
        },
    ]

    subjects: list[Subject] = []
    for spec in subject_catalog:
        subject = db.scalar(select(Subject).where(Subject.code == spec["code"]))
        if subject is None:
            subject = Subject(
                dept_id=dept.id,
                name=spec["name"],
                code=spec["code"],
                semester=spec["semester"],
                credits=4,
                max_marks=100,
            )
            db.add(subject)
            db.flush()
        subjects.append(subject)

    linked_subject_ids = set(
        db.scalars(
            select(TeacherSubject.subject_id).where(TeacherSubject.teacher_id == teacher.id)
        )
    )
    for subject in subjects:
        if subject.id not in linked_subject_ids:
            db.add(TeacherSubject(teacher_id=teacher.id, subject_id=subject.id))

    for subject, spec in zip(subjects, subject_catalog):
        existing_texts = set(
            db.scalars(select(Question.text).where(Question.subject_id == subject.id))
        )
        subject_questions: list[tuple[str, int, str, str, str, int]] = []
        for module_number, (module_title, focus_terms) in enumerate(
            spec["modules"], start=1
        ):
            subject_questions.extend(
                build_module_questions(
                    spec["name"],
                    module_title,
                    focus_terms,
                    module_number,
                )
            )

        for text, marks, co, bloom, difficulty, module in subject_questions:
            if text in existing_texts:
                continue
            existing_texts.add(text)
            db.add(
                Question(
                    subject_id=subject.id,
                    teacher_id=teacher.id,
                    text=text,
                    marks=marks,
                    course_outcome=co,
                    bloom_level=bloom,
                    difficulty=difficulty,
                    module_number=module,
                    tags=["seed", "dsatm", "aiml"],
                    is_verified=True,
                )
            )

    db.commit()
