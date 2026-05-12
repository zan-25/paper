from __future__ import annotations

import base64
import io
import json
import logging
import math
import re
import time
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import uuid4

import httpx
from docx import Document as DocxDocument
from fastapi import UploadFile
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.orm import Session

from .config import settings
from .generator import build_question_blueprint
from .models import Document, Question

logger = logging.getLogger("app.ai_service")

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}
TEXT_EXTENSIONS = {".txt", ".md", ".csv"}
RBT_LEVELS = ("L1", "L2", "L3", "L4", "L5", "L6")
CO_LEVELS = ("CO1", "CO2", "CO3", "CO4", "CO5", "CO6")
MODULE_RANGE = range(1, 6)
STOPWORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "for",
    "from",
    "how",
    "in",
    "into",
    "is",
    "it",
    "of",
    "on",
    "or",
    "that",
    "the",
    "to",
    "what",
    "which",
    "with",
}

EXTRACTION_SYSTEM_PROMPT = """You are an expert academic parser for VTU engineering question banks.

TASK:
- Read the uploaded material carefully.
- Extract every complete question.
- Infer metadata even when it is not explicitly written.
- Return ONLY valid JSON matching the required schema.

OUTPUT:
{
  "questions": [
    {
      "text": "Complete question text",
      "marks": 5,
      "bloom_level": "L1",
      "difficulty": "easy",
      "module_number": 1,
      "course_outcome": "CO1",
      "confidence": 0.0
    }
  ]
}

RULES:
- Bloom levels must be one of L1, L2, L3, L4, L5, L6.
- Course outcomes must be one of CO1, CO2, CO3, CO4, CO5, CO6.
- Difficulty must be easy, medium, or hard.
- Module number must be between 1 and 5 where possible.
- Preserve the original question wording as closely as possible.
"""


@dataclass
class ExtractedQuestion:
    text: str
    marks: int
    bloom_level: str
    difficulty: str
    module_number: int
    course_outcome: str
    confidence: float
    tags: list[str]


@dataclass
class ProcessingResult:
    success: bool
    document_id: int | None
    filename: str
    questions: list[ExtractedQuestion]
    total_extracted: int
    auto_approved: int
    processing_time: float
    ai_model: str
    ai_mode: str
    summary: dict[str, Any]
    error: str | None = None


@dataclass
class SelectionResult:
    questions: list[Question]
    coverage_stats: dict[str, Any]


@dataclass
class QuestionBankSummary:
    total_documents: int
    total_questions: int
    verified_questions: int
    pending_questions: int
    retrieval_ready_questions: int
    by_module: dict[str, int]
    by_rbt: dict[str, int]
    by_co: dict[str, int]
    by_difficulty: dict[str, int]
    recent_documents: list[dict[str, Any]]
    gaps: list[str]


class OllamaClient:
    def __init__(self) -> None:
        self.base_url = settings.ollama_base_url.rstrip("/")
        self.model = settings.ollama_model
        self.timeout = 60.0

    async def generate(
        self, prompt: str, system: str, images: list[str] | None = None
    ) -> dict[str, Any] | None:
        payload: dict[str, Any] = {
            "model": self.model,
            "prompt": prompt,
            "system": system,
            "stream": False,
            "format": "json",
        }
        if images:
            payload["images"] = images

        try:
            async with httpx.AsyncClient(timeout=self.timeout) as client:
                response = await client.post(
                    f"{self.base_url}/api/generate", json=payload
                )
                response.raise_for_status()
                raw = response.json().get("response", "{}")
                return json.loads(raw)
        except Exception as exc:
            logger.warning("Ollama request failed: %s", exc)
            return None

    async def is_available(self) -> bool:
        try:
            async with httpx.AsyncClient(timeout=5.0) as client:
                response = await client.get(f"{self.base_url}/api/tags")
                return response.status_code == 200
        except Exception:
            return False


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _tokenize(text: str) -> list[str]:
    tokens = re.findall(r"[a-z0-9]+", text.lower())
    return [token for token in tokens if len(token) > 2 and token not in STOPWORDS]


def build_sparse_vector(text: str) -> dict[str, float]:
    counts = Counter(_tokenize(text))
    if not counts:
        return {}
    length = math.sqrt(sum(value * value for value in counts.values()))
    return {token: round(value / length, 6) for token, value in counts.items()}


def vector_similarity(left: dict[str, float], right: dict[str, float]) -> float:
    if not left or not right:
        return 0.0
    if len(left) > len(right):
        left, right = right, left
    return sum(value * right.get(token, 0.0) for token, value in left.items())


def _extract_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return content.decode("utf-8", errors="ignore")
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".docx":
        document = DocxDocument(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text
        ]
        return "\n".join(paragraphs + table_cells)
    return content.decode("utf-8", errors="ignore")


def _guess_bloom_level(text: str) -> str:
    lowered = text.lower()
    rules = {
        "L1": ("define", "list", "state", "name", "identify", "recall"),
        "L2": ("explain", "describe", "discuss", "summarize", "outline"),
        "L3": ("solve", "calculate", "apply", "demonstrate", "implement"),
        "L4": ("analyze", "compare", "distinguish", "differentiate", "examine"),
        "L5": ("evaluate", "justify", "critique", "assess", "argue"),
        "L6": ("design", "develop", "construct", "create", "formulate"),
    }
    for level, keywords in rules.items():
        if any(keyword in lowered for keyword in keywords):
            return level
    return "L2"


def _guess_marks(text: str) -> int:
    length = len(text)
    if length < 80:
        return 5
    if length < 180:
        return 10
    return 15


def _guess_difficulty(marks: int, bloom_level: str) -> str:
    if marks <= 5 and bloom_level in {"L1", "L2"}:
        return "easy"
    if marks >= 15 or bloom_level in {"L5", "L6"}:
        return "hard"
    return "medium"


def _infer_course_outcome(text: str, bloom_level: str) -> str:
    lowered = text.lower()
    if bloom_level == "L1":
        return "CO1"
    if bloom_level in {"L5", "L6"}:
        return "CO5"
    if any(token in lowered for token in ("analyze", "comparison", "classification")):
        return "CO3"
    if any(token in lowered for token in ("design", "architecture", "pipeline")):
        return "CO4"
    return "CO2"


def heuristic_extract(text: str) -> list[ExtractedQuestion]:
    questions: list[ExtractedQuestion] = []
    current_module = 1

    for raw_line in re.split(r"\n+", text):
        line = _normalize_text(raw_line)
        if len(line) < 15:
            continue

        module_match = re.search(r"\bmodule\s*([1-5])\b", line, re.IGNORECASE)
        if module_match:
            current_module = int(module_match.group(1))
            continue

        looks_like_question = (
            bool(re.match(r"^\d+[\).\s-]", line))
            or "?" in line
            or any(
                keyword in line.lower()
                for keyword in (
                    "define",
                    "explain",
                    "describe",
                    "compare",
                    "analyze",
                    "design",
                    "solve",
                    "justify",
                )
            )
        )
        if not looks_like_question:
            continue

        clean = re.sub(r"^[\dA-Za-z]+[\).\s-]*", "", line).strip()
        clean = _normalize_text(clean or line)
        if len(clean) < 15:
            continue

        bloom_level = _guess_bloom_level(clean)
        marks = _guess_marks(clean)
        questions.append(
            ExtractedQuestion(
                text=clean,
                marks=marks,
                bloom_level=bloom_level,
                difficulty=_guess_difficulty(marks, bloom_level),
                module_number=current_module,
                course_outcome=_infer_course_outcome(clean, bloom_level),
                confidence=0.58,
                tags=["heuristic", "retrieval-ready"],
            )
        )

    return questions[:120]


def _parse_llm_questions(payload: dict[str, Any] | None) -> list[ExtractedQuestion]:
    if not isinstance(payload, dict):
        return []
    raw_questions = payload.get("questions")
    if not isinstance(raw_questions, list):
        return []

    parsed: list[ExtractedQuestion] = []
    for item in raw_questions:
        if not isinstance(item, dict):
            continue
        text = _normalize_text(str(item.get("text", "")))
        if len(text) < 10:
            continue
        bloom_level = str(item.get("bloom_level", _guess_bloom_level(text))).upper()
        if bloom_level not in RBT_LEVELS:
            bloom_level = _guess_bloom_level(text)
        marks = item.get("marks", _guess_marks(text))
        if isinstance(marks, str):
            match = re.search(r"\d+", marks)
            marks = int(match.group(0)) if match else _guess_marks(text)
        marks = max(5, min(int(marks), 20))
        difficulty = str(
            item.get("difficulty", _guess_difficulty(marks, bloom_level))
        ).lower()
        if difficulty not in {"easy", "medium", "hard"}:
            difficulty = _guess_difficulty(marks, bloom_level)
        module_number = int(item.get("module_number", 1) or 1)
        if module_number not in MODULE_RANGE:
            module_number = 1
        course_outcome = str(
            item.get("course_outcome", _infer_course_outcome(text, bloom_level))
        ).upper()
        if course_outcome not in CO_LEVELS:
            course_outcome = _infer_course_outcome(text, bloom_level)
        confidence = float(item.get("confidence", 0.8) or 0.8)
        parsed.append(
            ExtractedQuestion(
                text=text,
                marks=marks,
                bloom_level=bloom_level,
                difficulty=difficulty,
                module_number=module_number,
                course_outcome=course_outcome,
                confidence=max(0.0, min(confidence, 1.0)),
                tags=["llm", "retrieval-ready"],
            )
        )
    return parsed


def _dedupe_questions(questions: list[ExtractedQuestion]) -> list[ExtractedQuestion]:
    deduped: list[ExtractedQuestion] = []
    seen: set[str] = set()
    for question in questions:
        key = _normalize_text(question.text).lower()
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(question)
    return deduped


def _summarize_questions(questions: list[ExtractedQuestion]) -> dict[str, Any]:
    by_module = {str(module): 0 for module in MODULE_RANGE}
    by_rbt = {level: 0 for level in RBT_LEVELS}
    by_co = {level: 0 for level in CO_LEVELS}
    by_difficulty = {"easy": 0, "medium": 0, "hard": 0}

    for question in questions:
        by_module[str(question.module_number)] = (
            by_module.get(str(question.module_number), 0) + 1
        )
        by_rbt[question.bloom_level] = by_rbt.get(question.bloom_level, 0) + 1
        by_co[question.course_outcome] = by_co.get(question.course_outcome, 0) + 1
        by_difficulty[question.difficulty] = (
            by_difficulty.get(question.difficulty, 0) + 1
        )

    return {
        "by_module": by_module,
        "by_rbt": by_rbt,
        "by_co": by_co,
        "by_difficulty": by_difficulty,
    }


async def process_question_bank(
    file: UploadFile, subject_id: int, teacher_id: int, db: Session
) -> ProcessingResult:
    start_time = time.time()
    client = OllamaClient()

    content = await file.read()
    filename = file.filename or "upload"
    suffix = Path(filename).suffix.lower()
    is_image = suffix in IMAGE_EXTENSIONS

    parsed_text = ""
    if not is_image:
        parsed_text = _extract_text(filename, content)[:50000]

    ai_mode = "heuristic"
    extracted: list[ExtractedQuestion] = []

    if is_image:
        encoded = base64.b64encode(content).decode("utf-8")
        vision_prompt = (
            "Read this question bank image carefully and extract all academic questions."
        )
        extracted = _parse_llm_questions(
            await client.generate(
                vision_prompt,
                EXTRACTION_SYSTEM_PROMPT,
                images=[encoded],
            )
        )
        ai_mode = "vision" if extracted else "heuristic"
    elif parsed_text.strip():
        prompt = (
            "Extract academic questions from the following uploaded source material.\n\n"
            f"{parsed_text[:16000]}"
        )
        extracted = _parse_llm_questions(
            await client.generate(prompt, EXTRACTION_SYSTEM_PROMPT)
        )
        ai_mode = "llm-text" if extracted else "heuristic"

    if not extracted:
        extracted = heuristic_extract(parsed_text or filename)

    extracted = _dedupe_questions(extracted)
    if not extracted:
        return ProcessingResult(
            success=False,
            document_id=None,
            filename=filename,
            questions=[],
            total_extracted=0,
            auto_approved=0,
            processing_time=time.time() - start_time,
            ai_model=client.model,
            ai_mode=ai_mode,
            summary={},
            error="No questions could be extracted from the uploaded file",
        )

    upload_dir = settings.storage_path / "uploads" / str(subject_id)
    upload_dir.mkdir(parents=True, exist_ok=True)
    saved_path = upload_dir / f"{uuid4().hex}_{filename}"
    saved_path.write_bytes(content)

    document = Document(
        subject_id=subject_id,
        teacher_id=teacher_id,
        filename=filename,
        mime_type=file.content_type or "application/octet-stream",
        storage_path=str(saved_path),
        parsed_text=parsed_text or "\n".join(question.text for question in extracted)[:50000],
        upload_status="processed",
    )
    db.add(document)
    db.flush()

    saved_questions: list[Question] = []
    for question in extracted:
        tags = sorted(set(question.tags + [f"source:{suffix or 'raw'}", "vectorized"]))
        row = Question(
            subject_id=subject_id,
            teacher_id=teacher_id,
            source_doc_id=document.id,
            text=question.text,
            marks=question.marks,
            course_outcome=question.course_outcome,
            bloom_level=question.bloom_level,
            difficulty=question.difficulty,
            module_number=question.module_number,
            tags=tags,
            is_verified=question.confidence >= 0.82,
        )
        db.add(row)
        saved_questions.append(row)

    db.commit()

    summary = _summarize_questions(extracted)
    auto_approved = sum(1 for question in extracted if question.confidence >= 0.82)

    return ProcessingResult(
        success=True,
        document_id=document.id,
        filename=filename,
        questions=extracted,
        total_extracted=len(extracted),
        auto_approved=auto_approved,
        processing_time=time.time() - start_time,
        ai_model=client.model,
        ai_mode=ai_mode,
        summary=summary,
    )


def _distribution_targets(distribution: dict[str, int], slots: int) -> dict[str, int]:
    cleaned = {str(key).upper(): max(int(value), 0) for key, value in distribution.items() if int(value) > 0}
    if not cleaned:
        return {}
    total = sum(cleaned.values())
    if total <= 0:
        return {}

    exact = {key: (value / total) * slots for key, value in cleaned.items()}
    targets = {key: int(math.floor(value)) for key, value in exact.items()}
    remaining = max(0, slots - sum(targets.values()))
    remainders = sorted(
        ((value - targets[key], key) for key, value in exact.items()),
        reverse=True,
    )
    for _, key in remainders:
        if remaining <= 0:
            break
        targets[key] += 1
        remaining -= 1
    return targets


def _difficulty_score(desired: str, actual: str) -> float:
    desired = (desired or "balanced").lower()
    actual = actual.lower()
    if desired in {"", "balanced", "mixed"}:
        return 0.45
    if desired == actual:
        return 1.25
    if desired == "medium" and actual in {"easy", "hard"}:
        return 0.35
    return 0.1


def _question_selection_score(
    question: Question,
    query_vector: dict[str, float],
    rbt_remaining: dict[str, int],
    co_remaining: dict[str, int],
    uncovered_modules: set[int],
    target_module: int | None,
    desired_difficulty: str,
    target_marks: int,
) -> float:
    metadata_text = (
        f"{question.text} {question.course_outcome} {question.bloom_level} "
        f"module {question.module_number} {question.difficulty}"
    )
    score = vector_similarity(query_vector, build_sparse_vector(metadata_text)) * 7.0

    if rbt_remaining.get(question.bloom_level, 0) > 0:
        score += 3.0
    if co_remaining.get(question.course_outcome.upper(), 0) > 0:
        score += 2.4
    if question.module_number in uncovered_modules:
        score += 1.8
    if target_module is not None:
        if question.module_number == target_module:
            score += 3.5
        else:
            score -= 1.6

    marks_gap = abs((question.marks or target_marks) - target_marks)
    score += max(0.0, 1.2 - (marks_gap / max(target_marks, 1)))
    score += _difficulty_score(desired_difficulty, question.difficulty)
    score += 0.25 if question.is_verified else 0.0
    score += min(len(question.tags or []), 4) * 0.05
    return score


def _build_generation_prompt(
    subject_id: int,
    max_marks: int,
    modules: list[int],
    difficulty: str,
    rbt_distribution: dict[str, int],
    co_coverage: dict[str, int],
    blueprint: list[dict[str, Any]],
    candidates: list[Question],
    prompt: str | None,
) -> str:
    candidate_lines = []
    for question in candidates:
        candidate_lines.append(
            {
                "id": question.id,
                "text": question.text[:280],
                "module": question.module_number,
                "marks": question.marks,
                "co": question.course_outcome,
                "rbt": question.bloom_level,
                "difficulty": question.difficulty,
                "verified": question.is_verified,
            }
        )

    return json.dumps(
        {
            "task": "Select questions for a DSATM question paper template.",
            "subject_id": subject_id,
            "teacher_prompt": prompt or "",
            "requirements": {
                "max_marks": max_marks,
                "modules": modules,
                "difficulty": difficulty,
                "rbt_distribution": rbt_distribution,
                "co_coverage": co_coverage,
            },
            "template_slots": [
                {
                    "slot_label": slot["label"],
                    "question_number": slot["question_number"],
                    "subpart": slot["subpart"],
                    "slot_marks": slot["marks"],
                    "module_number": slot["module_number"],
                }
                for slot in blueprint
            ],
            "candidate_questions": candidate_lines,
            "response_schema": {
                "assignments": [
                    {
                        "slot_label": "1a",
                        "question_id": 123,
                    }
                ]
            },
        },
        ensure_ascii=True,
    )


def _parse_ai_assignments(
    payload: dict[str, Any] | None,
    blueprint: list[dict[str, Any]],
    allowed_ids: set[int],
) -> list[int]:
    if not isinstance(payload, dict):
        return []
    assignments = payload.get("assignments")
    if not isinstance(assignments, list):
        return []

    slot_order = [slot["label"] for slot in blueprint]
    assignment_map: dict[str, int] = {}
    used_ids: set[int] = set()
    for assignment in assignments:
        if not isinstance(assignment, dict):
            continue
        slot_label = str(assignment.get("slot_label", "")).strip()
        question_id = assignment.get("question_id")
        if (
            slot_label not in slot_order
            or not isinstance(question_id, int)
            or question_id not in allowed_ids
            or question_id in used_ids
        ):
            continue
        assignment_map[slot_label] = question_id
        used_ids.add(question_id)

    return [assignment_map[label] for label in slot_order if label in assignment_map]


def _calculate_coverage_stats(
    questions: list[Question],
    slot_marks: list[int],
    requested_modules: list[int],
    requested_rbt: dict[str, int],
    requested_co: dict[str, int],
) -> dict[str, Any]:
    by_module = {str(module): 0 for module in (requested_modules or list(MODULE_RANGE))}
    by_rbt = {level: 0 for level in RBT_LEVELS}
    by_co = {level: 0 for level in CO_LEVELS}

    for question, marks in zip(questions, slot_marks):
        by_module[str(question.module_number)] = (
            by_module.get(str(question.module_number), 0) + marks
        )
        by_rbt[question.bloom_level] = by_rbt.get(question.bloom_level, 0) + marks
        by_co[question.course_outcome.upper()] = (
            by_co.get(question.course_outcome.upper(), 0) + marks
        )

    total = sum(slot_marks) or 1
    return {
        "question_count": len(questions),
        "by_module": by_module,
        "by_rbt": by_rbt,
        "by_co": by_co,
        "requested": {
            "modules": requested_modules,
            "rbt": requested_rbt,
            "co": requested_co,
        },
        "percentages": {
            "co": {
                key: round((value / total) * 100)
                for key, value in by_co.items()
                if key in requested_co or value
            },
            "modules": {
                key: round((value / total) * 100)
                for key, value in by_module.items()
                if int(key) in requested_modules or value
            },
        },
    }


async def select_questions_for_paper(
    db: Session,
    subject_id: int,
    max_marks: int,
    modules: list[int],
    rbt_distribution: dict[str, int],
    co_coverage: dict[str, int],
    difficulty: str,
    prompt: str | None = None,
) -> SelectionResult:
    blueprint = build_question_blueprint(max_marks)
    slot_marks = [int(slot["marks"]) for slot in blueprint]
    client = OllamaClient()
    stmt = select(Question).where(Question.subject_id == subject_id)
    if modules:
        stmt = stmt.where(Question.module_number.in_(modules))
    candidates = list(db.scalars(stmt))
    if not candidates:
        return SelectionResult(
            questions=[],
            coverage_stats=_calculate_coverage_stats(
                [], [], modules, rbt_distribution, co_coverage
            ),
        )

    slot_count = len(blueprint)
    requested_modules = modules or sorted({question.module_number for question in candidates})
    rbt_targets = _distribution_targets(rbt_distribution, slot_count)
    co_targets = _distribution_targets(co_coverage, slot_count)
    query = prompt or (
        f"Generate question paper for subject {subject_id} with modules {requested_modules}, "
        f"RBT {rbt_distribution}, CO {co_coverage}, difficulty {difficulty}, max marks {max_marks}"
    )
    query_vector = build_sparse_vector(query)

    remaining = candidates[:]
    uncovered_modules = set(requested_modules)
    heuristic_selected: list[Question] = []
    ranked_candidates: list[Question] = []

    for slot in blueprint:
        if not remaining or len(heuristic_selected) >= min(slot_count, len(candidates)):
            break
        target_module = (
            int(slot["module_number"]) if slot.get("module_number") is not None else None
        )
        slot_candidates = (
            [question for question in remaining if question.module_number == target_module]
            if target_module is not None
            else remaining
        )
        candidate_pool = slot_candidates or remaining
        scored_remaining = sorted(
            candidate_pool,
            key=lambda question: _question_selection_score(
                question,
                query_vector,
                rbt_targets,
                co_targets,
                uncovered_modules,
                target_module,
                difficulty,
                int(slot["marks"]),
            ),
            reverse=True,
        )
        ranked_candidates.extend(scored_remaining[:3])
        best = scored_remaining[0]
        heuristic_selected.append(best)
        remaining.remove(best)
        uncovered_modules.discard(best.module_number)
        if rbt_targets.get(best.bloom_level, 0) > 0:
            rbt_targets[best.bloom_level] -= 1
        best_co = best.course_outcome.upper()
        if co_targets.get(best_co, 0) > 0:
            co_targets[best_co] -= 1

    selected = heuristic_selected
    if ranked_candidates and await client.is_available():
        unique_candidates: dict[int, Question] = {
            question.id: question for question in ranked_candidates
        }
        for question in heuristic_selected:
            unique_candidates[question.id] = question

        ai_prompt = _build_generation_prompt(
            subject_id,
            max_marks,
            requested_modules,
            difficulty,
            rbt_distribution,
            co_coverage,
            blueprint,
            list(unique_candidates.values())[: min(len(unique_candidates), 80)],
            prompt,
        )
        ai_result = await client.generate(
            ai_prompt,
            "You are the central academic planning brain for DSATM. Return only valid JSON.",
        )
        ai_question_ids = _parse_ai_assignments(
            ai_result,
            blueprint,
            set(unique_candidates.keys()),
        )
        if ai_question_ids:
            question_by_id = {question.id: question for question in candidates}
            ai_selected = [
                question_by_id[question_id]
                for question_id in ai_question_ids
                if question_id in question_by_id
            ]
            if len(ai_selected) >= max(6, min(slot_count, len(heuristic_selected)) // 2):
                selected = ai_selected

    assigned_slot_marks = slot_marks[: len(selected)]
    coverage = _calculate_coverage_stats(
        selected,
        assigned_slot_marks,
        requested_modules,
        rbt_distribution,
        co_coverage,
    )
    return SelectionResult(questions=selected[:slot_count], coverage_stats=coverage)


def summarize_question_bank(
    db: Session,
    subject_ids: list[int] | None = None,
    teacher_id: int | None = None,
) -> QuestionBankSummary:
    question_stmt = select(Question)
    document_stmt = select(Document)
    if subject_ids:
        question_stmt = question_stmt.where(Question.subject_id.in_(subject_ids))
        document_stmt = document_stmt.where(Document.subject_id.in_(subject_ids))
    if teacher_id is not None:
        question_stmt = question_stmt.where(Question.teacher_id == teacher_id)
        document_stmt = document_stmt.where(Document.teacher_id == teacher_id)

    questions = list(db.scalars(question_stmt))
    documents = list(db.scalars(document_stmt.order_by(Document.created_at.desc())))

    by_module = {str(module): 0 for module in MODULE_RANGE}
    by_rbt = {level: 0 for level in RBT_LEVELS}
    by_co = {level: 0 for level in CO_LEVELS}
    by_difficulty = {"easy": 0, "medium": 0, "hard": 0}

    for question in questions:
        by_module[str(question.module_number)] = (
            by_module.get(str(question.module_number), 0) + 1
        )
        by_rbt[question.bloom_level] = by_rbt.get(question.bloom_level, 0) + 1
        by_co[question.course_outcome.upper()] = (
            by_co.get(question.course_outcome.upper(), 0) + 1
        )
        by_difficulty[question.difficulty.lower()] = (
            by_difficulty.get(question.difficulty.lower(), 0) + 1
        )

    doc_question_counts: dict[int, int] = defaultdict(int)
    for question in questions:
        if question.source_doc_id is not None:
            doc_question_counts[question.source_doc_id] += 1

    recent_documents = [
        {
            "id": document.id,
            "filename": document.filename,
            "upload_status": document.upload_status,
            "created_at": document.created_at,
            "question_count": doc_question_counts.get(document.id, 0),
        }
        for document in documents[:10]
    ]

    gaps: list[str] = []
    for module in MODULE_RANGE:
        if by_module[str(module)] == 0:
            gaps.append(f"Module {module} has no indexed questions yet.")
    for level in RBT_LEVELS:
        if by_rbt[level] == 0:
            gaps.append(f"{level} has no coverage in the current bank.")
    for level in CO_LEVELS[:5]:
        if by_co[level] == 0:
            gaps.append(f"{level} is missing from the indexed question bank.")

    return QuestionBankSummary(
        total_documents=len(documents),
        total_questions=len(questions),
        verified_questions=sum(1 for question in questions if question.is_verified),
        pending_questions=sum(1 for question in questions if not question.is_verified),
        retrieval_ready_questions=len(questions),
        by_module=by_module,
        by_rbt=by_rbt,
        by_co=by_co,
        by_difficulty=by_difficulty,
        recent_documents=recent_documents,
        gaps=gaps[:8],
    )
