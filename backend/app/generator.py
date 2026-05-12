from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


@dataclass
class PaperConfig:
    department: str
    subject: str
    subject_code: str
    semester: str
    max_marks: int
    duration: str
    date: str
    batch: str
    teaching_department: str
    exam_type: str
    modules: list[int]
    rbt_levels: list[str]
    co_targets: list[str]
    year: str = "2026"
    instructions: str = "Instruction: Answer the following questions"
    college_name: str = "Dayananda Sagar Academy of Technology & Management"
    affiliation: str = "(Autonomous Institute under VTU)"
    program_line: str = "6 Programs Accredited by NBA (CSE, ISE, ECE, EEE, MECH, CV)"
    co_descriptions: dict[str, str] = field(default_factory=dict)
    co_percentages: dict[str, int] = field(default_factory=dict)
    module_percentages: dict[str, int] = field(default_factory=dict)
    left_seal_label: str = "DSATM"
    right_seal_label: str = "IQAC"
    template_note: str | None = None
    template_family: str = "dsatm"


def build_question_blueprint(max_marks: int) -> list[dict[str, Any]]:
    blueprint: list[dict[str, Any]] = []
    if max_marks <= 50:
        patterns = [(5, 5)] * 4 + [(4, 6)] * 6
        for question_number, (part_a, part_b) in enumerate(patterns, start=1):
            module_number = ((question_number - 1) // 2) + 1
            blueprint.extend(
                [
                    {
                        "question_number": question_number,
                        "subpart": "a",
                        "label": f"{question_number}a",
                        "marks": part_a,
                        "module_number": module_number,
                    },
                    {
                        "question_number": question_number,
                        "subpart": "b",
                        "label": f"{question_number}b",
                        "marks": part_b,
                        "module_number": module_number,
                    },
                ]
            )
        return blueprint

    hundred_mark_modules = {
        1: [(1, [("a", 6), ("b", 6), ("c", 8)]), (2, [("a", 6), ("b", 6), ("c", 8)])],
        2: [(3, [("a", 5), ("b", 8), ("c", 7)]), (4, [("a", 5), ("b", 8), ("c", 7)])],
        3: [(5, [("a", 5), ("b", 8), ("c", 7)]), (6, [("a", 5), ("b", 8), ("c", 7)])],
        4: [(7, [("a", 10), ("b", 10)]), (8, [("a", 10), ("b", 10)])],
        5: [(9, [("a", 10), ("b", 10)]), (10, [("a", 10), ("b", 10)])],
    }
    for module_number, question_sets in hundred_mark_modules.items():
        for question_number, parts in question_sets:
            for subpart, marks in parts:
                blueprint.append(
                    {
                        "question_number": question_number,
                        "subpart": subpart,
                        "label": f"{question_number}{subpart}",
                        "marks": marks,
                        "module_number": module_number,
                    }
                )
    return blueprint


class DSATMQuestionPaperGenerator:
    def generate(
        self, config: PaperConfig, questions: list[dict[str, Any]]
    ) -> DocumentType:
        document = Document()
        self._set_page_layout(document)
        self._add_header(document, config)
        self._add_usn_row(document)
        self._add_department_heading(document, config)
        self._add_exam_title(document, config)
        self._add_meta_table(document, config)
        self._add_instruction(document, config)
        self._add_questions_table(document, config, questions)
        self._add_course_outcomes_table(document, config)
        self._add_coverage_page(document, config)
        return document

    def save(self, document: DocumentType, output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return output_path

    def _set_page_layout(self, document: DocumentType) -> None:
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
        section.start_type = WD_SECTION.NEW_PAGE

        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(9)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    def _style_run(
        self,
        run,
        *,
        size: int = 9,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor | None = None,
    ) -> None:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color

    def _set_table_borders(
        self,
        table,
        *,
        top: bool = True,
        left: bool = True,
        bottom: bool = True,
        right: bool = True,
        inside_h: bool = True,
        inside_v: bool = True,
        size: str = "10",
    ) -> None:
        borders = OxmlElement("w:tblBorders")
        mapping = {
            "top": top,
            "left": left,
            "bottom": bottom,
            "right": right,
            "insideH": inside_h,
            "insideV": inside_v,
        }
        for name, enabled in mapping.items():
            border = OxmlElement(f"w:{name}")
            border.set(qn("w:val"), "single" if enabled else "nil")
            border.set(qn("w:sz"), size)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            borders.append(border)

        table_element = table._tbl
        table_pr = table_element.tblPr
        existing = table_pr.first_child_found_in("w:tblBorders")
        if existing is not None:
            table_pr.remove(existing)
        table_pr.append(borders)

    def _set_cell(
        self,
        cell,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        size: int = 9,
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    ) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        self._style_run(run, size=size, bold=bold, italic=italic)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _add_header(self, document: DocumentType, config: PaperConfig) -> None:
        table = document.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(
            table,
            top=False,
            left=False,
            bottom=False,
            right=False,
            inside_h=False,
            inside_v=False,
        )

        widths = [0.8, 4.0, 1.7, 0.8]
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)

        left_cell, title_cell, approval_cell, right_cell = table.rows[0].cells
        self._set_seal(left_cell, config.left_seal_label)

        title = title_cell.paragraphs[0]
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run(config.college_name)
        self._style_run(run, size=11, bold=True)
        title.add_run("\n")
        aff = title.add_run(config.affiliation)
        self._style_run(aff, size=8)

        approval = approval_cell.paragraphs[0]
        approval.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._append_highlighted_line(
            approval,
            [("Affiliated to ", False), ("VTU", True)],
            size=7.5,
        )
        self._append_highlighted_line(
            approval,
            [("Approved by ", False), ("AICTE", True)],
            size=7.5,
        )
        self._append_highlighted_line(
            approval,
            [("Accredited by ", False), ("NAAC", True), (" with A+ Grade", True)],
            size=7.5,
        )
        self._append_highlighted_line(
            approval,
            [("6 Programs Accredited by ", False), ("NBA", True)],
            size=7.5,
        )
        line = approval.add_run("\n(CSE, ISE, ECE, EEE, MECH, CV)")
        self._style_run(line, size=7.2)

        self._set_seal(right_cell, config.right_seal_label)

        divider = document.add_paragraph()
        divider.paragraph_format.space_before = Pt(2)
        divider.paragraph_format.space_after = Pt(2)
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = divider.add_run("_" * 116)
        self._style_run(run, size=7)

    def _set_seal(self, cell, label: str) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        top = paragraph.add_run(label)
        self._style_run(top, size=10, bold=True)
        bottom = paragraph.add_run("\nSeal")
        self._style_run(bottom, size=7)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _append_highlighted_line(
        self,
        paragraph,
        parts: list[tuple[str, bool]],
        *,
        size: float,
    ) -> None:
        for index, (text, emphasized) in enumerate(parts):
            run = paragraph.add_run(text)
            color = RGBColor(198, 40, 40) if emphasized else None
            self._style_run(run, size=int(size), color=color)
            if index == len(parts) - 1:
                paragraph.add_run("\n")

    def _add_usn_row(self, document: DocumentType) -> None:
        table = document.add_table(rows=1, cols=11)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table.autofit = False
        self._set_table_borders(table, top=False, left=False, bottom=False, right=False)

        label_cell = table.rows[0].cells[0]
        label_cell.width = Inches(0.55)
        self._set_cell(
            label_cell,
            "USN:",
            size=8,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        for index in range(1, 11):
            table.columns[index].width = Inches(0.31)
            self._set_cell(
                table.rows[0].cells[index],
                "",
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            self._set_table_borders(
                table,
                top=False,
                left=False,
                bottom=False,
                right=False,
                inside_h=False,
                inside_v=False,
            )
            cell = table.rows[0].cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "8")
                border.set(qn("w:color"), "000000")
                borders.append(border)
            tc_pr.append(borders)

    def _add_department_heading(self, document: DocumentType, config: PaperConfig) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"Department of {config.department}")
        self._style_run(run, size=12, bold=True)

    def _add_exam_title(self, document: DocumentType, config: PaperConfig) -> None:
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        table.columns[0].width = Inches(7.35)
        self._set_cell(
            table.rows[0].cells[0],
            config.exam_type,
            bold=True,
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    def _add_meta_table(self, document: DocumentType, config: PaperConfig) -> None:
        table = document.add_table(rows=5, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        widths = [1.5, 3.0, 1.5, 1.35]
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)

        data_rows = [
            ("Subject:", config.subject, "Subject Code:", config.subject_code),
            ("Semester:", config.semester, "Max. Marks:", str(config.max_marks)),
            ("Batch:", config.batch, "Duration:", config.duration),
            ("Date of IAT:", config.date, "Teaching Department:", config.teaching_department),
        ]
        for row_index, row in enumerate(data_rows):
            self._set_cell(table.rows[row_index].cells[0], row[0], bold=True)
            self._set_cell(table.rows[row_index].cells[1], row[1])
            self._set_cell(table.rows[row_index].cells[2], row[2], bold=True)
            self._set_cell(table.rows[row_index].cells[3], row[3])

        self._set_cell(table.rows[4].cells[0], "RBT Levels:", bold=True)
        merged = table.rows[4].cells[1].merge(table.rows[4].cells[3])
        rbt_text = (
            "L1-Remember, L2-Understand, L3-Apply, "
            "L4-Analyze, L5-Evaluate, L6-Create"
        )
        self._set_cell(merged, rbt_text)

    def _add_instruction(self, document: DocumentType, config: PaperConfig) -> None:
        if config.template_note:
            note = document.add_paragraph()
            note.paragraph_format.space_before = Pt(4)
            note.paragraph_format.space_after = Pt(3)
            note_label = note.add_run("Note:")
            self._style_run(note_label, size=9, bold=True)

            note_body = document.add_paragraph()
            note_body.paragraph_format.space_before = Pt(0)
            note_body.paragraph_format.space_after = Pt(4)
            body = note_body.add_run(config.template_note)
            self._style_run(body, size=9)

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(config.instructions)
        self._style_run(run, size=9, italic=True)

    def _add_questions_table(
        self,
        document: DocumentType,
        config: PaperConfig,
        questions: list[dict[str, Any]],
    ) -> None:
        blueprint = build_question_blueprint(config.max_marks)
        extra_rows = 4
        if config.max_marks > 50:
            extra_rows += 6
        table = document.add_table(rows=1 + len(blueprint) + extra_rows, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        widths = [0.75, 4.55, 0.7, 0.55, 0.65]
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)

        headers = ["Q No", "Questions", "Marks", "COs", "RBTL"]
        for index, header in enumerate(headers):
            self._set_cell(
                table.rows[0].cells[index],
                header,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

        padded_questions = questions[: len(blueprint)] + [{} for _ in range(max(0, len(blueprint) - len(questions[: len(blueprint)])))]
        row_index = 1
        current_module: int | None = None
        for slot, question in zip(blueprint, padded_questions):
            if config.max_marks > 50 and slot["module_number"] != current_module:
                current_module = int(slot["module_number"])
                module_row = table.rows[row_index]
                merged = module_row.cells[0].merge(module_row.cells[4])
                self._set_cell(
                    merged,
                    f"Module - {current_module}",
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                row_index += 1

            if (
                config.max_marks > 50
                and slot["subpart"] == "a"
                and slot["question_number"] % 2 == 0
            ):
                or_row = table.rows[row_index]
                merged = or_row.cells[0].merge(or_row.cells[4])
                self._set_cell(
                    merged,
                    "OR",
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                row_index += 1

            current = table.rows[row_index]
            question_text = str(question.get("text", ""))
            question_marks = int(question.get("marks", slot["marks"]) or slot["marks"])
            rbt_value = str(question.get("bloom_level", ""))
            self._set_cell(current.cells[0], slot["label"], align=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell(current.cells[1], question_text)
            self._set_cell(current.cells[2], str(question_marks), align=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell(
                current.cells[3],
                str(question.get("course_outcome", "")),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            self._set_cell(
                current.cells[4],
                rbt_value,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            current.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            row_index += 1

    def _add_course_outcomes_table(
        self,
        document: DocumentType,
        config: PaperConfig,
    ) -> None:
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(2)
        label = heading.add_run(
            "Course Outcomes (COs):  At the end of the Course, the Student will be able to:"
        )
        self._style_run(label, size=8.5, bold=True)

        table = document.add_table(rows=5, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        table.columns[0].width = Inches(0.6)
        table.columns[1].width = Inches(6.7)
        for index in range(1, 6):
            co_key = f"CO{index}"
            self._set_cell(
                table.rows[index - 1].cells[0],
                co_key,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            self._set_cell(
                table.rows[index - 1].cells[1],
                config.co_descriptions.get(co_key, ""),
            )

    def _add_coverage_page(self, document: DocumentType, config: PaperConfig) -> None:
        document.add_page_break()
        co_heading = document.add_paragraph()
        run = co_heading.add_run("Percentage of CO Coverage")
        self._style_run(run, size=9, bold=True)

        co_table = document.add_table(rows=2, cols=6)
        co_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        co_table.autofit = False
        self._set_table_borders(co_table)
        for index in range(6):
            co_table.columns[index].width = Inches(1.2 if index else 1.4)
        self._set_cell(co_table.rows[0].cells[0], "Course Outcomes", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell(co_table.rows[1].cells[0], "Percentage", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for index in range(1, 6):
            co_key = f"CO{index}"
            self._set_cell(co_table.rows[0].cells[index], co_key, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell(
                co_table.rows[1].cells[index],
                str(config.co_percentages.get(co_key, 0)),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

        module_heading = document.add_paragraph()
        module_heading.paragraph_format.space_before = Pt(12)
        run = module_heading.add_run("Percentage of Syllabus coverage")
        self._style_run(run, size=9, bold=True)

        module_table = document.add_table(rows=2, cols=6)
        module_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        module_table.autofit = False
        self._set_table_borders(module_table)
        for index in range(6):
            module_table.columns[index].width = Inches(1.05 if index else 1.55)
        self._set_cell(module_table.rows[0].cells[0], "Modules Covered", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell(module_table.rows[1].cells[0], "Percentage", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for index in range(1, 6):
            module_key = index
            self._set_cell(module_table.rows[0].cells[index], str(module_key), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell(
                module_table.rows[1].cells[index],
                str(config.module_percentages.get(str(module_key), config.module_percentages.get(module_key, 0))),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )


def generate_question_paper(
    config: PaperConfig, questions: list[dict[str, Any]], output_dir: Path
) -> Path:
    generator = DSATMQuestionPaperGenerator()
    document = generator.generate(config, questions)
    filename = (
        f"QP_{config.subject_code}_{config.exam_type}_{config.date.replace('-', '')}.docx"
    )
    return generator.save(document, output_dir / filename)


docx_generator = DSATMQuestionPaperGenerator()
